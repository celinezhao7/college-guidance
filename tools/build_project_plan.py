from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "College_Guidance_项目策划书.docx"

NAVY = "243B68"
BLUE = "5E76B8"
LAVENDER = "F3F0FA"
PALE_BLUE = "EEF3FB"
PALE_ORANGE = "FFF4EA"
ORANGE = "D4864C"
INK = "34364A"
MUTED = "666A7B"
LINE = "DDE1EA"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        set_cell_text(cell, h, True, WHITE, 9.5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            shade(cells[i], WHITE if row_idx % 2 == 0 else "F8F9FC")
            set_cell_text(cells[i], str(value), False, INK, 9)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text, fill=LAVENDER, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.10)
    table.columns[1].width = Inches(6.25)
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    table.cell(0, 0).text = ""
    cell = table.cell(0, 1)
    set_cell_margins(cell, 150, 180, 150, 180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    for r2 in p2.runs:
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


_PAGE_BREAK_COUNT = 0


def page_break(doc):
    """Keep the cover separate; let all later sections flow naturally."""
    global _PAGE_BREAK_COUNT
    _PAGE_BREAK_COUNT += 1
    if _PAGE_BREAK_COUNT == 1:
        doc.add_page_break()


def add_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.text = "COLLEGE GUIDANCE  ·  项目策划书"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("内部工作文档  |  v1.0  |  ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.3
    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 14, 8),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, NAVY, 8, 4),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = style_name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(65)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COLLEGE GUIDANCE")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    title = doc.add_paragraph("项目策划书", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("双语升学探索与申请内容推荐平台", style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_callout(
        doc,
        "文档用途",
        "供团队快速理解项目，也作为后续产品迭代、技术维护与决策记录的共同基线。本文描述当前已实现状态与建议路线图。",
        PALE_BLUE,
        ORANGE,
    )
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    for i, (k, v) in enumerate((
        ("版本", "v1.0"), ("日期", "2026 年 8 月 18 日"),
        ("项目负责人", "待填写"), ("文档状态", "内部评审稿"),
    )):
        set_cell_text(meta.cell(i, 0), k, True, NAVY, 9.5)
        set_cell_text(meta.cell(i, 1), v, False, INK, 9.5)
        shade(meta.cell(i, 0), LAVENDER)
        shade(meta.cell(i, 1), WHITE)
        meta.cell(i, 0).width = Inches(1.5)
        meta.cell(i, 1).width = Inches(4.7)
        set_cell_margins(meta.cell(i, 0), 110, 150, 110, 150)
        set_cell_margins(meta.cell(i, 1), 110, 150, 110, 150)

    page_break(doc)
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(
        "College Guidance 是一个面向大学申请探索阶段的双语 Web 产品。它把学生经历、官方申请指导、美国教育部公开院校数据与大语言模型结合起来，通过结构化对话帮助用户探索大学与专业领域，并选择更适合其经历的 Common App 或 UC PIQ 题目。"
    )
    add_callout(
        doc,
        "一句话定位",
        "不是一个泛化聊天机器人，而是一位有明确任务边界、能引用学生证据并逐步收集偏好的大学申请探索伙伴。",
        LAVENDER,
        BLUE,
    )
    doc.add_heading("当前状态", level=2)
    add_table(doc, ["维度", "已具备", "下一步重点"], [
        ("产品", "三项核心功能、学生档案、双语界面、快捷回答", "强化结果保存、比较与回访"),
        ("智能", "RAG、语义意图识别、院校名称校验、流式输出", "建立评测集与可解释性标准"),
        ("工程", "React/Vite 前端、FastAPI 后端、Docker、Render Blueprint", "监控、错误追踪、部署检查清单"),
        ("治理", "官方数据优先、已说明数据限制", "隐私政策、权限与数据保留规则"),
    ], [0.9, 2.6, 2.8])
    doc.add_heading("需要团队尽快确认的决策", level=2)
    for x in [
        "首要用户：独立申请学生、校内升学顾问，还是内部研究使用？",
        "是否允许用户创建账户并长期保存档案、对话与推荐结果？",
        "首个正式版本的成功标准与发布日期。",
        "模型调用、College Scorecard 与 Render 的月度成本上限及负责人。",
    ]:
        bullet(doc, x)

    page_break(doc)
    doc.add_heading("1. 项目背景与目标", level=1)
    doc.add_heading("1.1 问题定义", level=2)
    doc.add_paragraph(
        "大学申请信息分散、专业名称复杂，学生也常常难以把自身经历转化为选校或文书题目选择依据。普通聊天工具虽然灵活，但容易追问失焦、忽略数据约束，或给出缺少来源和边界的建议。"
    )
    doc.add_heading("1.2 产品目标", level=2)
    for x in [
        "把模糊目标转化为可回答、可推进的结构化对话。",
        "基于学生已提供的经历给出有证据的大学、专业领域或题目建议。",
        "在中文与英文环境中保持完整、一致的交互体验。",
        "使用官方来源和明确限制，降低用户把建议误解为确定结论的风险。",
    ]:
        bullet(doc, x)
    doc.add_heading("1.3 非目标", level=2)
    for x in [
        "不预测个人录取概率，也不承诺申请结果。",
        "不替代学校官网、招生办公室、专业顾问或法律/签证意见。",
        "不把 College Scorecard 的宽泛学科分类表述为学校课程目录中的精确专业。",
        "不自动代写整篇申请文书；重点是题目选择、素材发现与思路组织。",
    ]:
        bullet(doc, x)
    doc.add_heading("1.4 用户与场景", level=2)
    add_table(doc, ["用户", "典型需求", "产品回应"], [
        ("目标尚不清晰的学生", "不知道大学或专业从何开始", "用快捷回答与逐步追问建立偏好"),
        ("已有目标大学的学生", "想了解相关专业领域", "先核验学校，再结合经历探索领域"),
        ("准备申请文书的学生", "不知道哪道题最能展示自己", "从学生证据中推荐 Common App/PIQ 题目"),
        ("协作同事/顾问", "想快速理解推荐依据与系统边界", "提供结构化结果和清晰限制"),
    ], [1.4, 2.25, 2.65])

    page_break(doc)
    doc.add_heading("2. 产品方案", level=1)
    doc.add_heading("2.1 三项核心功能", level=2)
    add_table(doc, ["功能", "输入", "输出"], [
        ("大学与专业领域匹配", "目标大学/专业、地区、学校类型、规模、竞争程度等", "候选院校或专业领域及匹配理由"),
        ("Common App 文书题目推荐", "学生档案、经历与补充问题", "3 道更适合展开的题目及素材依据"),
        ("UC PIQ 题目推荐", "学生档案、经历与补充问题", "4 道更适合展示个人特质的题目及依据"),
    ], [1.65, 2.25, 2.4])
    doc.add_heading("2.2 关键体验设计", level=2)
    for title, text in [
        ("任务先于聊天", "首页直接呈现三项功能按钮，减少中央留白，也让产品不像普通 chatbot。"),
        ("快捷回答可选也可取消", "用结构化 choice_id 传递稳定语义；用户仍可自由输入，已选按钮可以取消。"),
        ("双语完整性", "语言切换后，按钮、占位符、加载状态、错误提示与生成状态应全部使用同一语言。"),
        ("渐进式收集", "每次只问一个关键问题；当用户说“不知道”时转向探索，而不是把它误当成学校名称。"),
        ("结果可解释", "建议应指出使用了哪些学生经历、偏好或官方数据，并同时说明限制。"),
    ]:
        doc.add_heading(title, level=3)
        doc.add_paragraph(text)
    doc.add_heading("2.3 核心用户流程", level=2)
    for x in [
        "选择学生档案、语言和一项功能；也可以取消功能选择后重新决定。",
        "系统根据模式进入对应对话；大学探索支持“先大学、先专业、两者都不确定”三种路径。",
        "快捷回答与自由文本共同收集偏好，语义分类器识别目标大学意图。",
        "当必要信息齐备后，系统调用检索与模型生成，并以流式方式展示结果。",
        "用户继续追问、调整偏好，未来可保存、比较或导出结果。",
    ]:
        bullet(doc, x)

    page_break(doc)
    doc.add_heading("3. 推荐逻辑与数据", level=1)
    doc.add_heading("3.1 工作机制", level=2)
    add_table(doc, ["阶段", "系统行为", "质量控制"], [
        ("理解", "识别模式、语言、choice_id 与自由文本意图", "结构化分类；低置信度时继续澄清"),
        ("验证", "对用户声称的目标大学进行 College Scorecard 核验", "核验成功后才写入会话状态"),
        ("检索", "从学生档案与官方指导的 Chroma 索引中检索相关片段", "优先使用与当前问题有关的证据"),
        ("生成", "Qwen 结合偏好、证据与数据生成建议", "提示词约束格式、边界与语言"),
        ("呈现", "流式返回推荐或结构化聊天响应", "本地/线上错误提示与加载状态本地化"),
    ], [0.8, 3.0, 2.5])
    doc.add_heading("3.2 数据来源", level=2)
    for x in [
        "美国教育部 College Scorecard：院校事实、学校类型、地点及宽泛 CIP 学科领域。",
        "UC 官方 PIQ 指导材料：用于理解题目目标与选择原则。",
        "Common App 官方指导材料：用于理解主文书题目与写作意图。",
        "学生档案与用户对话：用于个性化，不应被视为外部事实来源。",
    ]:
        bullet(doc, x)
    add_callout(
        doc,
        "数据边界",
        "College Scorecard 中的专业领域是联邦统计口径，不等同于学校官网的精确专业名称；目标大学可能因为名称、数据缺失或筛选条件而无法确认。建议结果必须保留这层说明。",
        PALE_ORANGE,
        ORANGE,
    )
    doc.add_heading("3.3 失败与回退原则", level=2)
    add_table(doc, ["情况", "应有行为"], [
        ("用户回答含糊", "不猜测；给出 2–4 个自然选项并允许自由输入"),
        ("院校无法核验", "提示使用英文官方名称，或转为“帮我探索”路径"),
        ("模型/API 暂时不可用", "保留用户输入，给出可恢复提示，不把技术错误伪装成推荐"),
        ("检索证据不足", "明确证据不足，提出补充问题或降低结论强度"),
    ], [1.75, 4.55])

    page_break(doc)
    doc.add_heading("4. 技术架构与运行", level=1)
    doc.add_heading("4.1 架构概览", level=2)
    add_table(doc, ["界面层", "服务层", "智能与数据层", "部署层"], [[
        "React 19\nTypeScript\nVite 8\nTailwind CSS 4",
        "FastAPI\n/api/chat\n/api/recommend\nprofiles / modes / health",
        "Qwen (DashScope)\nLangChain\nChroma\nCollege Scorecard API",
        "Docker\nRender Blueprint\ncollege-guidance",
    ]], [1.55, 1.65, 1.75, 1.35])
    doc.add_heading("4.2 接口与职责", level=2)
    add_table(doc, ["接口", "职责"], [
        ("GET /api/health", "Render 健康检查与服务状态"),
        ("GET /api/profiles", "返回可选择学生档案，不暴露本地路径"),
        ("GET /api/modes", "返回三项推荐模式及中英文标题"),
        ("POST /api/chat", "维护大学探索对话状态并返回快捷回答"),
        ("POST /api/recommend", "校验请求并流式生成推荐内容"),
    ], [1.75, 4.55])
    doc.add_heading("4.3 部署与配置", level=2)
    doc.add_paragraph(
        "Render 使用仓库根目录 Dockerfile 构建，服务名为 college-guidance，健康检查路径为 /api/health。敏感密钥不进入代码仓库，由 Render 环境变量管理。"
    )
    add_table(doc, ["变量", "用途", "管理建议"], [
        ("DASHSCOPE_API_KEY", "Qwen 调用凭据", "仅保存在部署平台；定期轮换"),
        ("DASHSCOPE_BASE_URL", "DashScope OpenAI 兼容地址", "配置化，不硬编码到业务逻辑"),
        ("QWEN_MODEL", "当前模型（qwen3.5-plus）", "升级前做回归评测"),
        ("COLLEGE_SCORECARD_API_KEY", "院校数据 API 凭据", "限制访问并监控配额"),
        ("COLLEGE_GUIDANCE_DEBUG", "调试开关", "生产环境保持 false"),
    ], [1.8, 2.2, 2.3])

    page_break(doc)
    doc.add_heading("5. 设计、内容与品牌原则", level=1)
    doc.add_heading("5.1 视觉方向", level=2)
    doc.add_paragraph(
        "当前界面保持简约框架：主内容区域以白色为主，左侧栏使用柔和粉紫蓝渐变，功能按钮之间形成连续的颜色过渡，并以极少量橙色制造温度和辨识度。橙色应是需要观察才会发现的点缀，而不是第四个主色。"
    )
    add_table(doc, ["元素", "原则"], [
        ("主背景", "白色，保证内容清晰并减少“AI 光效”感"),
        ("侧栏", "柔和梦幻渐变，但控制饱和度与对比度"),
        ("三项功能按钮", "同一色板中跨按钮连续过渡，不让单个按钮内部喧宾夺主"),
        ("橙色点缀", "用于小面积高光、状态或视觉平衡，不用于大块背景"),
        ("产品名", "左上角稳定展示 College Guidance，建立工具身份"),
    ], [1.55, 4.75])
    doc.add_heading("5.2 内容语气", level=2)
    for x in [
        "友好但不夸大：使用“帮助探索、提供建议”，避免“保证、最佳、一定录取”。",
        "一问一答：追问简短，说明为什么需要这条信息。",
        "语言一致：中文模式下不出现 Generating recommendation… 等英文状态文本。",
        "术语可理解：首次出现 PIQ、CIP、RAG 时给出中文解释。",
    ]:
        bullet(doc, x)
    doc.add_heading("5.3 无障碍与响应式", level=2)
    for x in [
        "按钮选中状态不能只依赖颜色；应同时提供边框、图标或 aria-pressed。",
        "快捷回答、发送按钮与错误信息需要键盘可达和清晰焦点状态。",
        "在移动端将侧栏改为顶部设置入口，保持输入框和功能选择优先可见。",
        "正文和浅色渐变背景保持足够对比度，避免目前偏淡颜色影响可读性。",
    ]:
        bullet(doc, x)

    page_break(doc)
    doc.add_heading("6. 质量、隐私与风险", level=1)
    doc.add_heading("6.1 测试策略", level=2)
    add_table(doc, ["层级", "重点", "当前/建议"], [
        ("单元测试", "意图分类、院校验证、快捷回答状态", "保留 test_conversation_service.py，并持续补充边界案例"),
        ("API 测试", "状态码、模式映射、流式响应、非法档案", "建立可重复的后端回归测试"),
        ("前端测试", "选中/取消、双语完整性、加载与失败恢复", "补充组件与端到端测试"),
        ("推荐评测", "相关性、证据一致、语言一致、边界遵守", "建立匿名化固定样本与人工评分表"),
        ("部署冒烟", "health、首页、chat、recommend", "每次上线后自动/人工检查"),
    ], [1.1, 2.6, 2.6])
    doc.add_heading("6.2 隐私与安全", level=2)
    for x in [
        "学生背景和偏好会发送至配置的 Qwen API；上线前需在界面与隐私说明中明确告知。",
        "仅收集完成推荐所需的信息，避免输入身份证件、财务或其他高敏感数据。",
        "学生档案、Chroma 索引、日志和导出文件需要定义保存期限、访问权限和删除流程。",
        "不得把模型输出当成事实数据库；院校事实应由官方来源验证。",
    ]:
        bullet(doc, x)
    doc.add_heading("6.3 主要风险与缓解", level=2)
    add_table(doc, ["风险", "影响", "缓解措施"], [
        ("模型幻觉或过度自信", "误导选校/选题", "检索证据、结构化输出、限制语、人工抽检"),
        ("官方数据口径过宽或缺失", "专业/院校信息不精确", "标注 CIP 边界，链接学校官网复核"),
        ("中英文体验不一致", "用户困惑、信任下降", "建立文案字典与双语回归检查"),
        ("第三方服务波动", "502、超时或生成失败", "健康检查、超时重试、可恢复错误提示"),
        ("学生信息泄露", "隐私与声誉风险", "最小化、权限、脱敏日志、密钥管理"),
        ("成本随使用增长", "预算不可控", "限流、缓存、使用指标与月度预算告警"),
    ], [1.55, 1.55, 3.2])

    page_break(doc)
    doc.add_heading("7. 指标与路线图", level=1)
    doc.add_heading("7.1 建议衡量指标", level=2)
    add_table(doc, ["目标", "指标示例", "说明"], [
        ("完成探索", "会话完成率、到达推荐结果比例", "判断追问是否过多或中途卡住"),
        ("建议有用", "结果收藏/复制/继续追问率、用户评分", "同时收集“不相关”的原因"),
        ("对话顺畅", "平均澄清次数、快捷回答使用率", "不是越少越好，要兼顾准确度"),
        ("系统可靠", "成功率、P95 响应时间、502/超时率", "按本地与 Render 环境区分"),
        ("输出可信", "事实核验通过率、证据一致率", "用固定评测集定期回归"),
        ("成本可控", "单次完成会话的模型成本", "结合缓存、模型选择和长度控制"),
    ], [1.2, 2.75, 2.35])
    doc.add_heading("7.2 分阶段路线图", level=2)
    add_table(doc, ["阶段", "目标", "建议交付物", "退出标准"], [
        ("阶段 1\n可靠 MVP", "稳定核心路径", "错误恢复、双语检查、评测集、监控", "三项功能在 Render 可稳定完成"),
        ("阶段 2\n可持续使用", "让结果可复用", "保存/导出、院校比较、会话历史、反馈入口", "用户能返回并继续之前的探索"),
        ("阶段 3\n团队运营", "形成产品闭环", "账户权限、顾问视图、内容管理、分析面板", "权限、隐私和运营流程可审计"),
    ], [1.15, 1.45, 2.35, 1.35])
    add_callout(
        doc,
        "当前优先级建议",
        "先把“可靠性 + 可解释性 + 结果保存”做扎实，再扩展更多学校数据库或复杂代理功能。对升学建议产品而言，可信与可回访比功能数量更能形成长期价值。",
        PALE_ORANGE,
        ORANGE,
    )
    doc.add_heading("7.3 建议的下一个迭代", level=2)
    for x in [
        "加入“保存本次结果 / 导出 PDF”与简短反馈按钮。",
        "建立 20–30 条中英文对话回归样本，覆盖“不知道”、缩写、错别字与取消选择。",
        "为 Render 增加结构化日志、错误追踪和部署后冒烟检查。",
        "在结果中统一呈现“推荐理由 / 使用证据 / 需要复核的信息 / 下一步行动”。",
    ]:
        numbered(doc, x)

    page_break(doc)
    doc.add_heading("8. 协作与维护", level=1)
    doc.add_heading("8.1 建议角色分工", level=2)
    add_table(doc, ["角色", "职责", "负责人"], [
        ("产品负责人", "范围、优先级、成功指标与版本验收", "待填写"),
        ("产品/交互设计", "流程、双语文案、视觉系统与可用性测试", "待填写"),
        ("前端", "界面、状态管理、可访问性与端到端体验", "待填写"),
        ("后端/AI", "API、会话逻辑、RAG、评测与数据校验", "待填写"),
        ("内容/升学顾问", "申请指导准确性、题目解释与人工评审", "待填写"),
        ("运维/安全", "Render、密钥、监控、备份与事件响应", "待填写"),
    ], [1.4, 3.5, 1.4])
    doc.add_heading("8.2 版本验收清单", level=2)
    for x in [
        "三项功能均能从首页进入、取消并重新选择。",
        "中文模式下界面、加载、错误、快捷回答与生成结果均为中文；英文同理。",
        "“不知道”等自然回答不会被当成大学名称；无法核验时提供可继续的路径。",
        "推荐说明所用学生证据，院校事实可追溯，限制清晰可见。",
        "API 密钥未进入仓库，Render health check 正常，生产 debug 为 false。",
        "自动测试通过，部署后完成 chat 与 recommend 冒烟测试。",
        "隐私说明、数据保留与删除方式已确认并可向用户解释。",
    ]:
        bullet(doc, x)
    doc.add_heading("8.3 决策记录模板", level=2)
    add_table(doc, ["日期", "决策", "原因", "负责人", "后续影响"], [
        ("待填写", "", "", "", ""),
        ("待填写", "", "", "", ""),
        ("待填写", "", "", "", ""),
    ], [0.9, 1.35, 1.55, 1.05, 1.45])

    page_break(doc)
    doc.add_heading("附录 A：当前事实清单", level=1)
    add_table(doc, ["类别", "当前实现"], [
        ("产品名称", "College Guidance"),
        ("前端", "React 19、TypeScript、Vite 8、Tailwind CSS 4"),
        ("后端", "FastAPI；/api/health、profiles、modes、chat、recommend"),
        ("模型与检索", "DashScope Qwen、LangChain、Chroma、本地索引"),
        ("外部数据", "美国教育部 College Scorecard"),
        ("部署", "Docker + Render Blueprint；服务名 college-guidance"),
        ("测试", "backend/tests/test_conversation_service.py 应持续保留"),
    ], [1.45, 4.85])
    doc.add_heading("附录 B：术语", level=1)
    add_table(doc, ["术语", "说明"], [
        ("PIQ", "UC Personal Insight Questions，UC 申请中的个人洞察问题"),
        ("CIP", "美国教育部学科分类体系；此处多为宽泛领域，不等于校内精确专业"),
        ("RAG", "检索增强生成：先检索相关证据，再让模型据此生成"),
        ("快捷回答", "界面上的建议选项；后端使用稳定 choice_id，而非依赖显示文字"),
        ("流式输出", "生成内容逐步显示，降低等待感"),
    ], [1.35, 4.95])
    doc.add_heading("附录 C：待确认问题", level=1)
    for x in [
        "正式目标用户、负责人和发布日期是什么？",
        "是否需要账户、云端档案、多人协作与顾问权限？",
        "哪些学生信息允许保存，保存多久，用户如何删除？",
        "正式发布前需要哪些法律、隐私或学校政策审查？",
        "是否需要接入学校官网的精确专业目录，及其维护方式？",
        "项目预算、模型成本上限和服务级别目标是什么？",
    ]:
        bullet(doc, x)

    doc.core_properties.title = "College Guidance 项目策划书"
    doc.core_properties.subject = "产品、技术与运营规划"
    doc.core_properties.author = "College Guidance 项目组"
    doc.core_properties.comments = "内部评审稿"
    doc.save(OUT)
    print("Project plan DOCX created.")


if __name__ == "__main__":
    build()
