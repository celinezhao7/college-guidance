import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from backend.app.profile_information import build_structured_profile
from backend.app.profile_service import StudentProfile
from backend.app.schemas import ProfileAdditionRecord


class StructuredProfileTests(unittest.TestCase):
    def test_original_and_confirmed_evidence_are_merged_with_sources(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "student.docx"
            document = Document()
            document.add_paragraph("Student Profile Summary\nAcademic Interest:\n- Biology\nCore Themes:\n- Curiosity")
            document.add_paragraph("@@@")
            document.add_paragraph("Experience 1: Lab Work\nCategory:\nAcademic\nResponsibilities:\nThe student:\n- Ran a test\nImpact:\n- Shared the result\nReflection:\nThe student realized:\n- Learned patience\nThemes:\ncuriosity, persistence")
            document.save(path)
            profile = StudentProfile("abc", "Student", "Student", "Student", path.name)
            addition = ProfileAdditionRecord(
                id="addition-1", confirmed_at="2026-08-21T00:00:00+00:00",
                experience_number=1, experience_title="Lab Work", action="", outcome="The result was presented.", reflection="",
            )
            with patch("backend.app.profile_information.get_profile_directory", return_value=Path(directory)):
                result = build_structured_profile(profile, [addition])

        self.assertEqual(["Biology"], result.academic_interests)
        self.assertEqual("enriched", result.experiences[0].status)
        self.assertEqual("Ran a test", result.experiences[0].action)
        self.assertEqual("Learned patience", result.experiences[0].reflection)
        self.assertIn("Shared the result", result.experiences[0].outcome)
        self.assertIn("The result was presented.", result.experiences[0].outcome)
        self.assertEqual(["original_profile", "user_confirmed"], [source.kind for source in result.experiences[0].sources])
        self.assertNotIn("outcome", result.experiences[0].missing_fields)

    def test_unmatched_addition_becomes_user_confirmed_experience(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "student.docx"
            document = Document()
            document.add_paragraph("Student Profile Summary")
            document.save(path)
            profile = StudentProfile("abc", "Student", "Student", "Student", path.name)
            addition = ProfileAdditionRecord(
                id="addition-2", confirmed_at="2026-08-21T00:00:00+00:00",
                experience_title="New Project", action="Built a prototype", outcome="", reflection="",
            )
            with patch("backend.app.profile_information.get_profile_directory", return_value=Path(directory)):
                result = build_structured_profile(profile, [addition])

        self.assertEqual("user_confirmed", result.experiences[0].status)
        self.assertEqual("New Project", result.experiences[0].experience_title)
        self.assertEqual(["outcome", "reflection"], result.experiences[0].missing_fields)


if __name__ == "__main__":
    unittest.main()
