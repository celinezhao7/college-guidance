import os
import hashlib
from pathlib import Path

from docx import Document as DocxDocument
from dotenv import load_dotenv

from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma

if __package__:
    from .student_profiles import get_student_profile_dir, list_student_profiles
else:
    from student_profiles import get_student_profile_dir, list_student_profiles


load_dotenv()


# -------------------------
# Paths
# -------------------------

BASE_DIR = Path(__file__).resolve().parent.parent

UC_DIR = BASE_DIR / "data" / "uc_official"
COMMON_APP_DIR = BASE_DIR / "data" / "common_app_official"

UC_DB_DIR = BASE_DIR / "chroma" / "uc"
COMMON_APP_DB_DIR = BASE_DIR / "chroma" / "common_app"
STUDENT_DB_DIR = BASE_DIR / "chroma" / "student"


# -------------------------
# Load .docx text
# -------------------------

def load_docx(path: Path) -> str:
    doc = DocxDocument(path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    ]

    return "\n".join(paragraphs)


# -------------------------
# Split using @@@
# -------------------------

def split_by_separator(text: str):
    return [
        chunk.strip()
        for chunk in text.split("@@@")
        if chunk.strip()
    ]


# -------------------------
# Classify student chunks
# -------------------------

def classify_student_chunk(chunk: str) -> str:
    lowered = chunk.lower()

    if "student profile summary" in lowered:
        return "profile_summary"

    if "retrieval instructions" in lowered:
        return "retrieval_instruction"

    return "experience"


def stable_document_id(document: Document) -> str:
    identity = "|".join(
        [
            str(document.metadata.get("type", "")),
            str(document.metadata.get("source", "")),
            str(document.metadata.get("chunk_index", "")),
            str(document.metadata.get("chunk_role", "")),
        ]
    )
    return hashlib.sha256(identity.encode("utf-8")).hexdigest()


def rebuild_collection(
    documents: list[Document],
    embeddings: OpenAIEmbeddings,
    collection_name: str,
    persist_directory: Path,
) -> None:
    vectorstore = Chroma(
        collection_name=collection_name,
        persist_directory=str(persist_directory),
        embedding_function=embeddings,
    )
    vectorstore.reset_collection()
    if documents:
        vectorstore.add_documents(
            documents=documents,
            ids=[stable_document_id(document) for document in documents],
        )


# -------------------------
# Load UC documents
# -------------------------

def load_uc_documents():
    documents = []

    for path in UC_DIR.glob("*.docx"):
        text = load_docx(path)
        chunks = split_by_separator(text)

        for index, chunk in enumerate(chunks):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": path.name,
                        "type": "uc_official",
                        "application_type": "uc",
                        "chunk_index": index,
                    },
                )
            )

    return documents


# -------------------------
# Load Common App documents
# -------------------------

def load_common_app_documents():
    documents = []

    for path in COMMON_APP_DIR.glob("*.docx"):
        text = load_docx(path)
        chunks = split_by_separator(text)

        for index, chunk in enumerate(chunks):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": path.name,
                        "type": "common_app_official",
                        "application_type": "common_app",
                        "chunk_index": index,
                    },
                )
            )

    return documents


# -------------------------
# Load student documents
# -------------------------

def load_student_documents(profile_path: Path):
    documents = []

    text = load_docx(profile_path)
    chunks = split_by_separator(text)

    for index, chunk in enumerate(chunks):
        documents.append(
            Document(
                page_content=chunk,
                metadata={
                    "source": profile_path.name,
                    "type": "student_evidence",
                    "chunk_index": index,
                    "chunk_role": classify_student_chunk(chunk),
                },
            )
        )

    return documents


# -------------------------
# Build vector databases
# -------------------------

def main():

    student_profiles = list_student_profiles()
    if not student_profiles:
        raise FileNotFoundError(
            f"No .docx student profiles found in {get_student_profile_dir()}"
        )

    embeddings = OpenAIEmbeddings(
        model="text-embedding-v4",
        api_key=os.getenv("DASHSCOPE_API_KEY"),
        base_url=os.getenv("DASHSCOPE_BASE_URL"),
        check_embedding_ctx_length=False,
        dimensions=1024,
        chunk_size=10,
    )

    # -------------------------
    # Load chunks
    # -------------------------

    uc_chunks = load_uc_documents()
    common_app_chunks = load_common_app_documents()
    student_chunks = []
    for student_profile in student_profiles:
        student_chunks.extend(load_student_documents(student_profile))

    print(f"UC chunks: {len(uc_chunks)}")
    print(f"Common App chunks: {len(common_app_chunks)}")
    print(f"Student profiles: {len(student_profiles)}")
    print(f"Student chunks: {len(student_chunks)}")

    print("\nStudent chunk roles:")

    for doc in student_chunks:
        print(
            f"- {doc.metadata['source']} | {doc.metadata['chunk_role']}: "
            f"{doc.page_content.splitlines()[0]}"
        )

    # -------------------------
    # Build UC vector database
    # -------------------------

    rebuild_collection(
        documents=uc_chunks,
        embeddings=embeddings,
        collection_name="uc_official",
        persist_directory=str(UC_DB_DIR),
    )

    # -------------------------
    # Build Common App vector database
    # -------------------------

    rebuild_collection(
        documents=common_app_chunks,
        embeddings=embeddings,
        collection_name="common_app_official",
        persist_directory=str(COMMON_APP_DB_DIR),
    )

    # -------------------------
    # Build Student vector database
    # -------------------------

    rebuild_collection(
        documents=student_chunks,
        embeddings=embeddings,
        collection_name="student_evidence",
        persist_directory=str(STUDENT_DB_DIR),
    )

    # -------------------------
    # Done
    # -------------------------

    print("\nVector databases created successfully.")
    print(f"UC database: {UC_DB_DIR}")
    print(f"Common App database: {COMMON_APP_DB_DIR}")
    print(f"Student database: {STUDENT_DB_DIR}")


if __name__ == "__main__":
    main()
